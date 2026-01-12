"""
Summary display component with export functionality.
This widget shows the generated summary and provides options to copy or export it
in various formats (TXT, PDF, DOCX).
"""

import logging
from pathlib import Path
from typing import Optional

from PyQt6.QtCore import Qt, pyqtSignal
from PyQt6.QtGui import QClipboard, QTextOption
from PyQt6.QtWidgets import (
    QApplication,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QMessageBox,
    QPushButton,
    QSizePolicy,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    logger.warning("python-docx or reportlab not available. Export features may be limited.")


class SummaryDisplayWidget(QWidget):
    """Widget for displaying and exporting summaries."""
    
    refresh_clicked = pyqtSignal()  # Signal emitted when refresh button is clicked
    
    def __init__(self, parent=None):
        """Initialize summary display widget."""
        super().__init__(parent)
        self.summary_text = ""
        self.summarization_mode = ""  # Store the mode used for summarization
        self.length_preset = ""  # Store length preset for abstractive mode
        self._setup_ui()
    
    def _setup_ui(self):
        """
        Set up the UI components following HCI principles.
        - Enhanced typography for better readability
        - Card-based design with elevation
        - Clear action hierarchy
        """
        layout = QVBoxLayout()
        layout.setSpacing(12)  # 8px grid: 1.5 * 8 = 12px
        layout.setContentsMargins(0, 0, 0, 0)  # Consistent margins for alignment
        
        # Header with stats and actions - modern spacing
        header_layout = QHBoxLayout()
        header_layout.setSpacing(16)
        header_layout.setContentsMargins(0, 0, 0, 8)
        
        # Mode label with modern styling (shows which mode was used)
        self.mode_label = QLabel("")
        self.mode_label.setStyleSheet("""
            QLabel {
                color: #7c8ff5;
                font-size: 9pt;
                font-weight: 600;
                padding: 5px 12px;
                background-color: #1c1e26;
                border-radius: 6px;
                border: 1px solid #3a3d4a;
            }
        """)
        self.mode_label.hide()  # Hide initially until a summary is generated
        
        # Action buttons - modern with icons and better spacing
        button_layout = QHBoxLayout()
        button_layout.setSpacing(8)  # Better spacing
        
        # Modern button styling with subtle effects - compact size
        button_style = """
            QPushButton {
                font-size: 8.5pt;
                font-weight: 600;
                padding: 5px 10px;
                min-height: 28px;
                border-radius: 6px;
                letter-spacing: 0.2px;
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #1f212b, stop:1 #1a1c26);
                border: 1px solid #3a3d4a;
                color: #e8eaed;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #667eea, stop:1 #5568d3);
                border: 1px solid #7c8ff5;
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #4c52cc, stop:1 #3d42b8);
            }
            QPushButton:disabled {
                opacity: 0.4;
            }
        """
        
        self.copy_button = QPushButton("📋 Copy")
        self.copy_button.clicked.connect(self._copy_to_clipboard)
        self.copy_button.setEnabled(False)
        self.copy_button.setStyleSheet(button_style)
        
        self.export_txt_button = QPushButton("💾 TXT")
        self.export_txt_button.setToolTip("Export as Text File")
        self.export_txt_button.clicked.connect(self._export_txt)
        self.export_txt_button.setEnabled(False)
        self.export_txt_button.setStyleSheet(button_style)
        
        self.export_pdf_button = QPushButton("📄 PDF")
        self.export_pdf_button.setToolTip("Export as PDF")
        self.export_pdf_button.clicked.connect(self._export_pdf)
        self.export_pdf_button.setEnabled(False)
        self.export_pdf_button.setStyleSheet(button_style)
        
        self.export_docx_button = QPushButton("📝 DOCX")
        self.export_docx_button.setToolTip("Export as Word Document")
        self.export_docx_button.clicked.connect(self._export_docx)
        self.export_docx_button.setEnabled(False)
        self.export_docx_button.setStyleSheet(button_style)
        
        # Add refresh button to reset application
        self.refresh_button = QPushButton("🔄 Reset")
        self.refresh_button.setToolTip("Reset to initial state")
        self.refresh_button.clicked.connect(self._on_refresh)
        self.refresh_button.setEnabled(False)
        self.refresh_button.setStyleSheet(button_style)
        
        button_layout.addWidget(self.copy_button)
        button_layout.addWidget(self.export_txt_button)
        button_layout.addWidget(self.export_pdf_button)
        button_layout.addWidget(self.export_docx_button)
        button_layout.addWidget(self.refresh_button)
        
        header_layout.addWidget(self.mode_label)
        header_layout.addStretch()
        header_layout.addLayout(button_layout)
        
        # Text display with exceptional typography and card design
        self.text_display = QTextEdit()
        self.text_display.setReadOnly(True)
        self.text_display.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)  # Wrap at widget width
        self.text_display.setWordWrapMode(QTextOption.WrapMode.WordWrap)  # Wrap at word boundaries
        self.text_display.setAcceptRichText(False)  # Plain text only
        self.text_display.setStyleSheet("""
            QTextEdit {
                font-family: 'Segoe UI', 'SF Pro Display', 'Inter', 'Roboto', Arial, sans-serif;
                font-size: 10.5pt;
                line-height: 1.8;
                letter-spacing: 0.01em;
                padding: 20px 24px;
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #1f212b, stop:1 #1a1c26);
                border: 2px solid #3a3d4a;
                border-radius: 14px;
                color: #e8eaed;
            }
            QTextEdit:focus {
                border: 2px solid #7c8ff5;
            }
        """)
        
        layout.addLayout(header_layout)
        layout.addWidget(self.text_display, 1)  # Give text display stretch factor
        
        self.setLayout(layout)
        
        # Make text display expand to use available space
        self.text_display.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding
        )
        # Set minimum dimensions to ensure content is visible and uses space efficiently
        self.text_display.setMinimumHeight(200)
        self.text_display.setMinimumWidth(400)  # Ensure adequate width for text flow
    
    def set_summary(self, text: str, mode: str = "", length_preset: str = ""):
        """
        Set the summary text to display.
        
        Args:
            text: Summary text
            mode: Summarization mode used ("extractive" or "abstractive")
        """
        # Clean up text formatting for better display
        cleaned_text = self._clean_display_text(text)
        
        self.summary_text = cleaned_text
        self.summarization_mode = mode
        self.length_preset = (length_preset or "").lower()
        self.text_display.setPlainText(cleaned_text)
        self._update_mode_label()
        
        # Enable buttons
        has_text = bool(cleaned_text.strip())
        self.copy_button.setEnabled(has_text)
        self.export_txt_button.setEnabled(has_text)
        self.export_pdf_button.setEnabled(has_text)
        self.export_docx_button.setEnabled(has_text)
        self.refresh_button.setEnabled(has_text)
    
    def _clean_display_text(self, text: str) -> str:
        """
        Clean text for optimal display - fix spacing and line breaks.
        
        Args:
            text: Raw summary text
            
        Returns:
            Cleaned text optimized for display
        """
        if not text:
            return ""
        
        import re
        
        # Remove excessive whitespace but preserve intentional paragraph breaks
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Normalize line breaks - keep intentional paragraph breaks (double newlines)
        # but remove single newlines that break text flow
        text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 newlines (paragraph break)
        
        # For extractive mode: ensure sentences flow together properly
        # Remove single newlines within paragraphs, but keep double newlines
        if self.summarization_mode.lower() == "extractive":
            # Join lines that don't end with sentence-ending punctuation
            lines = text.split('\n')
            cleaned_lines = []
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    # Empty line - keep as paragraph break if not duplicate
                    if cleaned_lines and cleaned_lines[-1] != '':
                        cleaned_lines.append('')
                    continue
                
                # If previous line doesn't end with sentence terminator, join with space
                if cleaned_lines and cleaned_lines[-1] and not cleaned_lines[-1].rstrip()[-1:] in '.!?':
                    cleaned_lines[-1] = cleaned_lines[-1] + ' ' + line
                else:
                    cleaned_lines.append(line)
            
            text = '\n'.join(cleaned_lines)
        
        # Final cleanup: remove leading/trailing whitespace on each line
        lines = text.split('\n')
        lines = [line.strip() for line in lines]
        text = '\n'.join(lines)
        
        # Remove any leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_summary(self) -> str:
        """Get the current summary text."""
        return self.summary_text
    
    def clear_summary(self):
        """Clear the summary display."""
        self.summary_text = ""
        self.summarization_mode = ""
        self.length_preset = ""
        self.text_display.clear()
        self._update_mode_label()
        
        # Disable buttons
        self.copy_button.setEnabled(False)
        self.export_txt_button.setEnabled(False)
        self.export_pdf_button.setEnabled(False)
        self.export_docx_button.setEnabled(False)
        self.refresh_button.setEnabled(False)
    
    def _on_refresh(self):
        """Handle refresh button click - emit signal to main window."""
        self.refresh_clicked.emit()
    
    def _update_mode_label(self):
        """
        Update the mode label to show which summarization mode was used.
        """
        if not self.summarization_mode:
            self.mode_label.setText("")
            self.mode_label.hide()
        else:
            # Show the mode with appropriate icon
            if self.summarization_mode.lower() == "extractive":
                self.mode_label.setText("⚡ Extractive Mode")
            else:
                preset_suffix = ""
                if self.length_preset:
                    preset_suffix = f" • {self.length_preset.title()} preset"
                self.mode_label.setText(f"🤖 Abstractive Mode{preset_suffix}")
            self.mode_label.show()
    
    def _copy_to_clipboard(self):
        """
        Copy summary text to system clipboard.
        Shows a brief message if successful.
        """
        if not self.summary_text:
            logger.warning("Attempted to copy empty summary")
            return
        
        try:
            # Get the application's clipboard instance
            clipboard = QApplication.clipboard()
            clipboard.setText(self.summary_text)
            logger.info("Summary copied to clipboard successfully")
            
            # Show brief feedback (optional - could use a status bar message instead)
            # For now, we'll just log it
        except Exception as e:
            logger.error(f"Failed to copy to clipboard: {str(e)}")
            # Show error message to user
            QMessageBox.warning(
                self,
                "Copy Failed",
                f"Failed to copy summary to clipboard:\n{str(e)}"
            )
    
    def _export_txt(self):
        """
        Export summary to a plain text file.
        Opens a file dialog for the user to choose save location.
        """
        if not self.summary_text:
            logger.warning("Attempted to export empty summary")
            return
        
        try:
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Export Summary as TXT",
                "summary.txt",  # Default filename
                "Text Files (*.txt);;All Files (*)"
            )
            
            if file_path:
                # Ensure .txt extension if not provided
                if not file_path.endswith('.txt'):
                    file_path += '.txt'
                
                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(self.summary_text)
                    logger.info(f"Summary exported to TXT: {file_path}")
                except PermissionError:
                    error_msg = f"Permission denied. Cannot write to:\n{file_path}\n\nFile may be open in another program."
                    logger.error(error_msg)
                    QMessageBox.warning(self, "Export Failed", error_msg)
                except OSError as e:
                    error_msg = f"Failed to write file:\n{str(e)}"
                    logger.error(error_msg)
                    QMessageBox.warning(self, "Export Failed", error_msg)
                except Exception as e:
                    error_msg = f"Unexpected error exporting to TXT:\n{str(e)}"
                    logger.error(error_msg, exc_info=True)
                    QMessageBox.warning(self, "Export Failed", error_msg)
        except Exception as e:
            logger.error(f"Error in export dialog: {str(e)}")
            QMessageBox.warning(self, "Error", f"Failed to open export dialog:\n{str(e)}")
    
    def _export_pdf(self):
        """
        Export summary to a PDF file.
        Uses reportlab library to create a properly formatted PDF document.
        """
        if not self.summary_text:
            logger.warning("Attempted to export empty summary to PDF")
            return
        
        # Check if reportlab is available
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
        except ImportError:
            error_msg = (
                "PDF export requires the 'reportlab' library.\n\n"
                "Please install it with: pip install reportlab"
            )
            logger.error("reportlab not available for PDF export")
            QMessageBox.warning(self, "Export Unavailable", error_msg)
            return
        
        try:
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Export Summary as PDF",
                "summary.pdf",  # Default filename
                "PDF Files (*.pdf);;All Files (*)"
            )
            
            if file_path:
                # Ensure .pdf extension
                if not file_path.endswith('.pdf'):
                    file_path += '.pdf'
                
                try:
                    # Create PDF canvas with standard letter size
                    c = canvas.Canvas(file_path, pagesize=letter)
                    width, height = letter
                    
                    # Set up margins and formatting
                    margin = 50
                    top_margin = height - 50
                    bottom_margin = margin
                    max_width = width - 2 * margin
                    
                    # Title
                    c.setFont("Helvetica-Bold", 16)
                    c.drawString(margin, top_margin, "Summary")
                    
                    # Summary text with proper word wrapping
                    c.setFont("Helvetica", 11)
                    y = top_margin - 30
                    line_height = 14
                    
                    # Split text into words and wrap lines
                    words = self.summary_text.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        text_width = c.stringWidth(test_line, "Helvetica", 11)
                        
                        # If line is too long, draw current line and start new one
                        if text_width > max_width and current_line:
                            c.drawString(margin, y, current_line)
                            y -= line_height
                            
                            # Check if we need a new page
                            if y < bottom_margin:
                                c.showPage()
                                y = height - margin
                            
                            current_line = word
                        else:
                            current_line = test_line
                    
                    # Draw the last line
                    if current_line:
                        c.drawString(margin, y, current_line)
                    
                    # Save the PDF
                    c.save()
                    logger.info(f"Summary exported to PDF: {file_path}")
                    
                except PermissionError:
                    error_msg = f"Permission denied. Cannot write to:\n{file_path}\n\nFile may be open in another program."
                    logger.error(error_msg)
                    QMessageBox.warning(self, "Export Failed", error_msg)
                except Exception as e:
                    error_msg = f"Error creating PDF:\n{str(e)}"
                    logger.error(error_msg, exc_info=True)
                    QMessageBox.warning(self, "Export Failed", error_msg)
        except Exception as e:
            logger.error(f"Error in PDF export dialog: {str(e)}")
            QMessageBox.warning(self, "Error", f"Failed to export PDF:\n{str(e)}")
    
    def _export_docx(self):
        """
        Export summary to a Microsoft Word document (DOCX format).
        Uses python-docx library to create a properly formatted document.
        """
        if not self.summary_text:
            logger.warning("Attempted to export empty summary to DOCX")
            return
        
        # Check if python-docx is available
        try:
            from docx import Document
        except ImportError:
            error_msg = (
                "DOCX export requires the 'python-docx' library.\n\n"
                "Please install it with: pip install python-docx"
            )
            logger.error("python-docx not available for DOCX export")
            QMessageBox.warning(self, "Export Unavailable", error_msg)
            return
        
        try:
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Export Summary as DOCX",
                "summary.docx",  # Default filename
                "Word Documents (*.docx);;All Files (*)"
            )
            
            if file_path:
                # Ensure .docx extension
                if not file_path.endswith('.docx'):
                    file_path += '.docx'
                
                try:
                    # Create a new Word document
                    doc = Document()
                    
                    # Add title
                    doc.add_heading('Summary', 0)
                    
                    # Add the summary text as a paragraph
                    # The library handles formatting automatically
                    doc.add_paragraph(self.summary_text)
                    
                    # Save the document
                    doc.save(file_path)
                    logger.info(f"Summary exported to DOCX: {file_path}")
                    
                except PermissionError:
                    error_msg = f"Permission denied. Cannot write to:\n{file_path}\n\nFile may be open in another program."
                    logger.error(error_msg)
                    QMessageBox.warning(self, "Export Failed", error_msg)
                except Exception as e:
                    error_msg = f"Error creating DOCX file:\n{str(e)}"
                    logger.error(error_msg, exc_info=True)
                    QMessageBox.warning(self, "Export Failed", error_msg)
        except Exception as e:
            logger.error(f"Error in DOCX export dialog: {str(e)}")
            QMessageBox.warning(self, "Error", f"Failed to export DOCX:\n{str(e)}")

